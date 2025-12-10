"""
Data ingestion service for processing multiple input sources.
This module provides functionality to extract text content from various sources
including PDFs, images (with OCR), documents, databases and streams.
"""

from pathlib import Path
from typing import Dict, Any, List, Optional, Union, BinaryIO, Tuple, AsyncGenerator
import logging
import os
import json
import csv
import io
import asyncio
from datetime import datetime
import aiofiles
import PyPDF2
from PIL import Image
import docx
import pandas as pd
import motor.motor_asyncio
from fastapi import UploadFile
import tempfile

logger = logging.getLogger(__name__)

# Try to import optional dependencies
try:
    import pytesseract
    TESSERACT_AVAILABLE = True
except ImportError:
    logger.warning("pytesseract not available. OCR functionality will be disabled.")
    TESSERACT_AVAILABLE = False

try:
    import fitz  # PyMuPDF
    PYMUPDF_AVAILABLE = True
except ImportError:
    logger.warning("PyMuPDF not available. Enhanced PDF processing will be disabled.")
    PYMUPDF_AVAILABLE = False

class DataIngestionService:
    """Service for ingesting data from multiple sources"""
    
    def __init__(self, upload_dir: str = "temp_uploads"):
        self.upload_dir = Path(upload_dir)
        self.upload_dir.mkdir(exist_ok=True)
        
    async def ingest_file(self, file_path: Union[str, Path]) -> Dict[str, Any]:
        """
        Ingest a file from the filesystem and extract its text content
        
        Args:
            file_path: Path to the file to process
            
        Returns:
            Dict containing text content and metadata
        """
        file_path = Path(file_path)
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        file_ext = file_path.suffix.lower()
        
        # Process based on file type
        if file_ext in ['.pdf']:
            return await self._process_pdf(file_path)
        elif file_ext in ['.png', '.jpg', '.jpeg', '.tiff', '.bmp', '.gif']:
            return await self._process_image(file_path)
        elif file_ext in ['.docx']:
            return await self._process_docx(file_path)
        elif file_ext in ['.xlsx', '.xls']:
            return await self._process_excel(file_path)
        elif file_ext in ['.csv']:
            return await self._process_csv(file_path)
        elif file_ext in ['.json']:
            return await self._process_json(file_path)
        elif file_ext in ['.txt', '.md', '.py', '.js', '.html', '.xml']:
            return await self._process_text_file(file_path)
        else:
            # Try to process as text
            try:
                return await self._process_text_file(file_path)
            except UnicodeDecodeError:
                raise ValueError(f"Unsupported file type: {file_ext}")
    
    async def ingest_uploaded_file(self, file: UploadFile) -> Dict[str, Any]:
        """
        Process an uploaded file from FastAPI
        
        Args:
            file: FastAPI UploadFile object
            
        Returns:
            Dict containing text content and metadata
        """
        # Create a temporary file
        suffix = Path(file.filename).suffix
        with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
            # Read uploaded content in chunks to avoid memory issues with large files
            content = await file.read()
            tmp.write(content)
            tmp_path = tmp.name
        
        try:
            # Process the temp file
            result = await self.ingest_file(tmp_path)
            return result
        finally:
            # Clean up the temporary file
            if os.path.exists(tmp_path):
                os.unlink(tmp_path)
    
    async def ingest_from_database(self, 
                                  connection_string: str,
                                  database_name: str,
                                  collection_name: str,
                                  query: Dict[str, Any],
                                  text_fields: List[str],
                                  limit: int = 100) -> List[Dict[str, Any]]:
        """
        Extract text data from a MongoDB database
        
        Args:
            connection_string: MongoDB connection string
            database_name: Name of the database
            collection_name: Name of the collection
            query: MongoDB query to filter documents
            text_fields: List of field names containing text to extract
            limit: Maximum number of documents to process
            
        Returns:
            List of dicts containing extracted text and metadata
        """
        client = motor.motor_asyncio.AsyncIOMotorClient(connection_string)
        db = client[database_name]
        collection = db[collection_name]
        
        results = []
        cursor = collection.find(query).limit(limit)
        
        async for document in cursor:
            # Extract specified text fields
            text_content = "\n\n".join(
                str(document.get(field, "")) 
                for field in text_fields 
                if field in document
            )
            
            # Create result with metadata
            result = {
                "text": text_content,
                "metadata": {
                    "source": "database",
                    "database": database_name,
                    "collection": collection_name,
                    "document_id": str(document.get("_id")),
                    "extracted_fields": [f for f in text_fields if f in document],
                    "timestamp": datetime.utcnow().isoformat()
                }
            }
            results.append(result)
        
        return results
    
    async def ingest_text(self, text: str, metadata: Dict[str, Any] = None) -> Dict[str, Any]:
        """
        Ingest raw text with optional metadata
        
        Args:
            text: Text content to process
            metadata: Optional metadata to include
            
        Returns:
            Dict containing text content and metadata
        """
        return {
            "text": text,
            "metadata": metadata or {
                "source": "direct_text",
                "timestamp": datetime.utcnow().isoformat()
            }
        }
    
    async def ingest_url(self, url: str, headers: Dict[str, str] = None) -> Dict[str, Any]:
        """
        Fetch and extract text content from a URL
        
        Args:
            url: URL to fetch
            headers: Optional HTTP headers
            
        Returns:
            Dict containing text content and metadata
        """
        try:
            import aiohttp
            
            async with aiohttp.ClientSession() as session:
                async with session.get(url, headers=headers) as response:
                    if response.status == 200:
                        content_type = response.headers.get('Content-Type', '')
                        
                        # Handle based on content type
                        if 'text/html' in content_type:
                            html_content = await response.text()
                            # Use a simple HTML text extraction (consider using BeautifulSoup in production)
                            import re
                            text_content = re.sub(r'<[^>]+>', ' ', html_content)
                            text_content = re.sub(r'\s+', ' ', text_content).strip()
                        elif 'application/json' in content_type:
                            json_data = await response.json()
                            # Convert JSON to string representation
                            text_content = json.dumps(json_data, indent=2)
                        elif 'text/' in content_type:
                            text_content = await response.text()
                        elif 'application/pdf' in content_type:
                            # Download PDF and process
                            pdf_content = await response.read()
                            with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as tmp:
                                tmp.write(pdf_content)
                                tmp_path = tmp.name
                            
                            try:
                                pdf_result = await self._process_pdf(tmp_path)
                                text_content = pdf_result["text"]
                            finally:
                                if os.path.exists(tmp_path):
                                    os.unlink(tmp_path)
                        else:
                            return {
                                "text": "",
                                "error": f"Unsupported content type: {content_type}",
                                "metadata": {
                                    "source": "url",
                                    "url": url,
                                    "timestamp": datetime.utcnow().isoformat(),
                                    "content_type": content_type
                                }
                            }
                        
                        return {
                            "text": text_content,
                            "metadata": {
                                "source": "url",
                                "url": url,
                                "timestamp": datetime.utcnow().isoformat(),
                                "content_type": content_type
                            }
                        }
                    else:
                        return {
                            "text": "",
                            "error": f"Failed to fetch URL: HTTP {response.status}",
                            "metadata": {
                                "source": "url",
                                "url": url,
                                "timestamp": datetime.utcnow().isoformat(),
                                "status_code": response.status
                            }
                        }
        
        except ImportError:
            raise ImportError("aiohttp package is required for URL ingestion")
    
    async def _process_pdf(self, file_path: Path) -> Dict[str, Any]:
        """Extract text from PDF file"""
        text_content = ""
        page_count = 0
        metadata = {}
        
        # Try PyMuPDF first if available (better results)
        if PYMUPDF_AVAILABLE:
            try:
                doc = fitz.open(file_path)
                page_count = len(doc)
                pages = []
                
                for page_num, page in enumerate(doc):
                    page_text = page.get_text()
                    pages.append(page_text)
                
                text_content = "\n\n".join(pages)
                
                # Extract document info
                metadata = doc.metadata
                
                doc.close()
            except Exception as e:
                logger.warning(f"PyMuPDF processing failed: {e}. Falling back to PyPDF2.")
                # Fall back to PyPDF2
        
        # Use PyPDF2 if PyMuPDF failed or is not available
        if not text_content:
            try:
                with open(file_path, "rb") as file:
                    reader = PyPDF2.PdfReader(file)
                    page_count = len(reader.pages)
                    pages = []
                    
                    for page_num in range(page_count):
                        page = reader.pages[page_num]
                        pages.append(page.extract_text() or "")
                    
                    text_content = "\n\n".join(pages)
                    
                    # Extract document info if available
                    if reader.metadata:
                        metadata = dict(reader.metadata)
            except Exception as e:
                raise ValueError(f"Failed to process PDF: {e}")
        
        return {
            "text": text_content,
            "metadata": {
                "source": "pdf",
                "file_path": str(file_path),
                "page_count": page_count,
                "document_info": metadata,
                "timestamp": datetime.utcnow().isoformat()
            }
        }
    
    async def _process_image(self, file_path: Path) -> Dict[str, Any]:
        """Extract text from image using OCR"""
        if not TESSERACT_AVAILABLE:
            raise ImportError("pytesseract not available. Install it for OCR support.")
        
        try:
            with Image.open(file_path) as img:
                # Get image metadata
                width, height = img.size
                format_name = img.format
                mode = img.mode
                
                # Perform OCR
                text_content = pytesseract.image_to_string(img)
                
                return {
                    "text": text_content,
                    "metadata": {
                        "source": "image",
                        "file_path": str(file_path),
                        "image_size": f"{width}x{height}",
                        "format": format_name,
                        "mode": mode,
                        "timestamp": datetime.utcnow().isoformat()
                    }
                }
                
        except Exception as e:
            raise ValueError(f"Failed to process image: {e}")
    
    async def _process_docx(self, file_path: Path) -> Dict[str, Any]:
        """Extract text from DOCX file"""
        try:
            doc = docx.Document(file_path)
            
            # Extract text from paragraphs
            text_content = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text for cell in row.cells if cell.text.strip())
                    if row_text:
                        text_content += f"\n{row_text}"
            
            # Extract document properties
            core_properties = {}
            try:
                props = doc.core_properties
                core_properties = {
                    "author": props.author,
                    "category": props.category,
                    "comments": props.comments,
                    "content_status": props.content_status,
                    "created": props.created.isoformat() if props.created else None,
                    "identifier": props.identifier,
                    "keywords": props.keywords,
                    "language": props.language,
                    "last_modified_by": props.last_modified_by,
                    "last_printed": props.last_printed.isoformat() if props.last_printed else None,
                    "modified": props.modified.isoformat() if props.modified else None,
                    "revision": props.revision,
                    "subject": props.subject,
                    "title": props.title,
                    "version": props.version
                }
                # Remove None values
                core_properties = {k: v for k, v in core_properties.items() if v is not None}
            except:
                pass
            
            return {
                "text": text_content,
                "metadata": {
                    "source": "docx",
                    "file_path": str(file_path),
                    "document_properties": core_properties,
                    "timestamp": datetime.utcnow().isoformat()
                }
            }
            
        except Exception as e:
            raise ValueError(f"Failed to process DOCX: {e}")
    
    async def _process_excel(self, file_path: Path) -> Dict[str, Any]:
        """Extract text from Excel file"""
        try:
            # Load workbook
            df_dict = pd.read_excel(file_path, sheet_name=None)
            
            text_content = []
            sheet_data = {}
            
            # Process each sheet
            for sheet_name, df in df_dict.items():
                # Convert dataframe to string with proper formatting
                sheet_text = f"Sheet: {sheet_name}\n"
                sheet_text += df.to_string(index=False)
                text_content.append(sheet_text)
                
                # Store structured sheet data
                sheet_data[sheet_name] = {
                    "columns": df.columns.tolist(),
                    "rows": len(df),
                    "sample": df.head(5).to_dict(orient="records")
                }
            
            return {
                "text": "\n\n".join(text_content),
                "metadata": {
                    "source": "excel",
                    "file_path": str(file_path),
                    "sheets": list(df_dict.keys()),
                    "sheet_data": sheet_data,
                    "timestamp": datetime.utcnow().isoformat()
                }
            }
            
        except Exception as e:
            raise ValueError(f"Failed to process Excel file: {e}")
    
    async def _process_csv(self, file_path: Path) -> Dict[str, Any]:
        """Extract text from CSV file"""
        try:
            df = pd.read_csv(file_path)
            
            # Generate text representation
            text_content = df.to_string(index=False)
            
            return {
                "text": text_content,
                "metadata": {
                    "source": "csv",
                    "file_path": str(file_path),
                    "columns": df.columns.tolist(),
                    "rows": len(df),
                    "sample": df.head(5).to_dict(orient="records"),
                    "timestamp": datetime.utcnow().isoformat()
                }
            }
            
        except Exception as e:
            # Try with different encoding if default fails
            try:
                df = pd.read_csv(file_path, encoding="latin1")
                text_content = df.to_string(index=False)
                
                return {
                    "text": text_content,
                    "metadata": {
                        "source": "csv",
                        "file_path": str(file_path),
                        "columns": df.columns.tolist(),
                        "rows": len(df),
                        "encoding": "latin1",
                        "timestamp": datetime.utcnow().isoformat()
                    }
                }
            except Exception as e2:
                raise ValueError(f"Failed to process CSV file: {e}, {e2}")
    
    async def _process_json(self, file_path: Path) -> Dict[str, Any]:
        """Extract text from JSON file"""
        try:
            async with aiofiles.open(file_path, "r", encoding="utf-8") as f:
                content = await f.read()
                data = json.loads(content)
                
                # Convert to formatted string
                text_content = json.dumps(data, indent=2)
                
                # Try to extract some basic structure info
                structure_info = {}
                if isinstance(data, dict):
                    structure_info["type"] = "object"
                    structure_info["keys"] = list(data.keys())
                elif isinstance(data, list):
                    structure_info["type"] = "array"
                    structure_info["length"] = len(data)
                    if data and isinstance(data[0], dict):
                        structure_info["sample_keys"] = list(data[0].keys())
                
                return {
                    "text": text_content,
                    "metadata": {
                        "source": "json",
                        "file_path": str(file_path),
                        "structure": structure_info,
                        "timestamp": datetime.utcnow().isoformat()
                    }
                }
                
        except Exception as e:
            raise ValueError(f"Failed to process JSON file: {e}")
    
    async def _process_text_file(self, file_path: Path) -> Dict[str, Any]:
        """Extract text from plain text file"""
        try:
            async with aiofiles.open(file_path, "r", encoding="utf-8") as f:
                text_content = await f.read()
                
                return {
                    "text": text_content,
                    "metadata": {
                        "source": "text",
                        "file_path": str(file_path),
                        "file_size": os.path.getsize(file_path),
                        "lines": text_content.count("\n") + 1,
                        "timestamp": datetime.utcnow().isoformat()
                    }
                }
        except UnicodeDecodeError:
            # Try with different encoding if UTF-8 fails
            try:
                async with aiofiles.open(file_path, "r", encoding="latin1") as f:
                    text_content = await f.read()
                    
                    return {
                        "text": text_content,
                        "metadata": {
                            "source": "text",
                            "file_path": str(file_path),
                            "encoding": "latin1",
                            "file_size": os.path.getsize(file_path),
                            "lines": text_content.count("\n") + 1,
                            "timestamp": datetime.utcnow().isoformat()
                        }
                    }
            except Exception as e2:
                raise ValueError(f"Failed to process text file: {e2}")
        except Exception as e:
            raise ValueError(f"Failed to process text file: {e}")
            
    async def process_stream(self, stream_generator, chunk_size: int = 1000) -> AsyncGenerator[Dict[str, Any], None]:
        """
        Process text from a stream or generator
        
        Args:
            stream_generator: Async generator that yields text chunks
            chunk_size: Maximum number of chunks to accumulate before processing
            
        Yields:
            Dict containing accumulated text content and metadata
        """
        buffer = []
        count = 0
        
        async for chunk in stream_generator:
            buffer.append(chunk)
            count += 1
            
            if count >= chunk_size:
                # Process the accumulated buffer
                text_content = "".join(buffer)
                yield {
                    "text": text_content,
                    "metadata": {
                        "source": "stream",
                        "chunk_count": count,
                        "timestamp": datetime.utcnow().isoformat()
                    }
                }
                # Reset the buffer
                buffer = []
                count = 0
        
        # Process any remaining content
        if buffer:
            text_content = "".join(buffer)
            yield {
                "text": text_content,
                "metadata": {
                    "source": "stream",
                    "chunk_count": count,
                    "timestamp": datetime.utcnow().isoformat()
                }
            }