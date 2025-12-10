"""
Document text extraction service with OCR and MRZ parsing capabilities.
Supports PDF, DOCX, XLSX, and image files.
"""
import os
import logging
import fitz  # PyMuPDF
import pdfplumber
import docx
import openpyxl
import pytesseract
import numpy as np
from PIL import Image, ImageEnhance
from passporteye import mrz
from passporteye.mrz.image import MRZPipeline
import io
import re
import tempfile
import cv2
from typing import Dict, List, Any, Optional, Tuple, Union

# Configure logging
logger = logging.getLogger(__name__)

# Supported file types
SUPPORTED_TYPES = ["pdf", "docx", "xlsx", "jpg", "jpeg", "png", "tiff", "tif", "bmp"]

# Configure tesseract path if needed (uncomment and modify as necessary)
# pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'

class MRZDetector:
    """Enhanced MRZ detection with multiple strategies for higher accuracy"""
    
    def __init__(self):
        self.mrz_patterns = {
            'td1': re.compile(r'[A-Z0-9<]{30}\n[A-Z0-9<]{30}\n[A-Z0-9<]{30}'),
            'td2': re.compile(r'[A-Z0-9<]{36}\n[A-Z0-9<]{36}'),
            'td3': re.compile(r'[A-Z0-9<]{44}\n[A-Z0-9<]{44}'),
            'mrva': re.compile(r'[A-Z0-9<]{44}\n[A-Z0-9<]{44}'),
            'mrvb': re.compile(r'[A-Z0-9<]{36}\n[A-Z0-9<]{36}')
        }
        self.pipeline = MRZPipeline(extra_cmdline_params='--oem 0')

    def detect_from_text(self, text: str) -> Optional[Dict[str, Any]]:
        """Detect MRZ from text using regex patterns"""
        text = text.replace(' ', '')
        
        # Try to match against known MRZ patterns
        for mrz_type, pattern in self.mrz_patterns.items():
            match = pattern.search(text)
            if match:
                try:
                    result = mrz.MRZ(match.group(0))
                    if result and result.valid:
                        return {
                            'mrz_type': mrz_type,
                            'valid': result.valid,
                            'valid_score': result.valid_score,
                            'data': result.to_dict(),
                        }
                except Exception as e:
                    logger.warning(f"MRZ parsing error: {e}")
        return None
    
    def detect_from_image(self, image: Image.Image) -> Optional[Dict[str, Any]]:
        """Detect MRZ from image using passporteye"""
        try:
            # Convert PIL Image to bytes
            with io.BytesIO() as output:
                # Enhance image for better MRZ detection
                image = self._preprocess_image(image)
                image.save(output, format='PNG')
                image_bytes = output.getvalue()
            
            # Process with passporteye
            result = self.pipeline.process(image_bytes=image_bytes)
            if result and result.valid:
                return {
                    'mrz_type': result.mrz_type,
                    'valid': result.valid,
                    'valid_score': result.valid_score,
                    'data': result.to_dict(),
                }
        except Exception as e:
            logger.warning(f"MRZ image processing error: {e}")
        return None
    
    def _preprocess_image(self, image: Image.Image) -> Image.Image:
        """Preprocess image to improve MRZ recognition"""
        # Convert to grayscale
        image = image.convert('L')
        
        # Enhance contrast
        enhancer = ImageEnhance.Contrast(image)
        image = enhancer.enhance(2.0)
        
        # Apply thresholding
        threshold = 150
        image = image.point(lambda p: p > threshold and 255)
        
        return image

class DocumentExtractor:
    """
    Extract text from various document types with OCR and MRZ detection
    """
    
    def __init__(self, ocr_config=None):
        """
        Initialize the extractor with optional OCR configuration
        """
        self.mrz_detector = MRZDetector()
        self.ocr_config = ocr_config or {}
        
    def _normalize_bbox(self, bbox, width=0, height=0, source='pdf') -> List[float]:
        """
        Normalize bounding box coordinates to a consistent format [x1, y1, x2, y2]
        """
        if not bbox:
            return None
        
        # PyMuPDF bbox is already [x1, y1, x2, y2]
        if source == 'pdf' and len(bbox) == 4:
            return [float(coord) for coord in bbox]
        
        # Tesseract bbox is [left, top, width, height]
        elif source == 'ocr' and len(bbox) == 4:
            x, y, w, h = bbox
            return [float(x), float(y), float(x + w), float(y + h)]
        
        return None

    def _is_page_scanned(self, page, text_length_threshold=100, text_area_ratio_threshold=0.01) -> bool:
        """
        Determine if a PDF page is scanned or has searchable text
        """
        # Extract text and analyze its length
        text = page.get_text("text")
        if len(text.strip()) < text_length_threshold:
            return True
            
        # Check text area coverage
        width, height = page.rect.width, page.rect.height
        page_area = width * height
        
        text_area = 0
        for b in page.get_text("dict")["blocks"]:
            if "lines" in b:
                for l in b["lines"]:
                    if "spans" in l:
                        for s in l["spans"]:
                            if "bbox" in s:
                                x0, y0, x1, y1 = s["bbox"]
                                text_area += (x1 - x0) * (y1 - y0)
        
        text_area_ratio = text_area / page_area
        return text_area_ratio < text_area_ratio_threshold
    
    def _apply_ocr(self, image: Image.Image, dpi=300) -> List[Dict[str, Any]]:
        """
        Apply OCR to an image and return structured text blocks
        """
        blocks = []
        try:
            # Preprocess image for better OCR
            image = self._preprocess_image_for_ocr(image)
            
            # Run OCR with tesseract
            custom_config = f'--oem 1 --psm 3 -l eng+fra+deu+spa'
            ocr_data = pytesseract.image_to_data(image, output_type=pytesseract.Output.DICT, config=custom_config)
            
            # Extract text blocks with confidence and bounding boxes
            for i in range(len(ocr_data["text"])):
                text = ocr_data["text"][i]
                if not text.strip():
                    continue
                    
                # Get bounding box
                left = ocr_data["left"][i]
                top = ocr_data["top"][i]
                width = ocr_data["width"][i]
                height = ocr_data["height"][i]
                bbox = [left, top, width, height]
                
                # Normalize confidence score to 0-1
                conf = float(ocr_data["conf"][i]) / 100 if ocr_data["conf"][i] != '-1' else 0.0
                
                # Set block type as OCR
                block_type = "ocr"
                
                # Check for MRZ in the text
                mrz_result = self.mrz_detector.detect_from_text(text)
                if mrz_result:
                    block_type = "mrz"
                
                # Add text block to results
                blocks.append({
                    "text": text,
                    "bbox": self._normalize_bbox(bbox, source='ocr'),
                    "conf": conf,
                    "type": block_type,
                    "metadata": {
                        "block_num": ocr_data["block_num"][i],
                        "line_num": ocr_data["line_num"][i],
                        "word_num": ocr_data["word_num"][i],
                        "mrz_data": mrz_result
                    }
                })
            
            # Check for MRZ in the image as a whole
            mrz_result = self.mrz_detector.detect_from_image(image)
            if mrz_result and mrz_result.get('valid', False):
                # Add MRZ block with the full image dimensions if detected
                blocks.append({
                    "text": str(mrz_result.get('data', {})),
                    "bbox": [0, 0, image.width, image.height],
                    "conf": mrz_result.get('valid_score', 0.8),
                    "type": "mrz",
                    "metadata": {
                        "mrz_type": mrz_result.get('mrz_type'),
                        "mrz_data": mrz_result
                    }
                })
                
        except Exception as e:
            logger.error(f"OCR processing error: {e}")
            
        return blocks
    
    def _preprocess_image_for_ocr(self, image: Image.Image) -> Image.Image:
        """
        Preprocess image to improve OCR quality
        """
        # Convert to grayscale if not already
        if image.mode != 'L':
            image = image.convert('L')
            
        # Resize if too small or too large
        dpi = 300
        if min(image.size) < 1000:
            # Calculate scaling factor to achieve about 300 DPI
            scale = dpi / 72
            new_size = (int(image.width * scale), int(image.height * scale))
            image = image.resize(new_size, Image.Resampling.LANCZOS)
            
        # Enhance contrast
        enhancer = ImageEnhance.Contrast(image)
        image = enhancer.enhance(1.5)
            
        return image

    def _handle_rotation(self, page) -> fitz.Page:
        """
        Handle page rotation for proper text extraction
        """
        # Check if page needs rotation
        rotation = page.rotation
        if rotation != 0:
            # Create a transformation matrix for rotation correction
            matrix = fitz.Matrix(1, 0, 0, 1, 0, 0).prerotate(360 - rotation)
            # Apply transformation
            page.apply_transform(matrix)
        
        return page
    
    def _extract_from_pdf(self, file_path: str) -> List[Dict[str, Any]]:
        """
        Extract text from PDF files, handling both normal and scanned PDFs
        """
        results = []
        try:
            # Open PDF with PyMuPDF
            doc = fitz.open(file_path)
            
            for page_num, page in enumerate(doc, 1):
                blocks = []
                page_info = {"width": page.rect.width, "height": page.rect.height, "rotation": page.rotation}
                
                # Apply rotation correction if needed
                page = self._handle_rotation(page)
                
                # Check if page appears to be scanned
                is_scanned = self._is_page_scanned(page)
                
                # Extract text blocks from PDF using PyMuPDF
                for b in page.get_text("blocks"):
                    text, bbox = b[4], b[:4]
                    conf = 1.0  # PyMuPDF doesn't provide confidence
                    block_type = "text"
                    
                    # Check for MRZ in text
                    mrz_result = self.mrz_detector.detect_from_text(text)
                    if mrz_result:
                        block_type = "mrz"
                        
                    # Add block to results
                    blocks.append({
                        "text": text,
                        "bbox": self._normalize_bbox(bbox, source='pdf'),
                        "conf": conf,
                        "type": block_type,
                        "metadata": {"mrz_data": mrz_result} if mrz_result else {}
                    })
                
                # Apply OCR for scanned pages or pages with little text
                if is_scanned:
                    # Render page to image
                    pix = page.get_pixmap(alpha=False)
                    img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                    
                    # Apply OCR
                    ocr_blocks = self._apply_ocr(img)
                    
                    # Add OCR blocks to results
                    blocks.extend(ocr_blocks)
                
                # Add page results
                results.append({
                    "page": page_num, 
                    "text_blocks": blocks,
                    "page_info": page_info,
                    "is_scanned": is_scanned
                })
                
        except Exception as e:
            logger.error(f"PDF processing error: {e}")
            raise
            
        return results
    
    def _extract_from_docx(self, file_path: str) -> List[Dict[str, Any]]:
        """
        Extract text from DOCX files
        """
        results = []
        try:
            # Open DOCX document
            doc = docx.Document(file_path)
            blocks = []
            
            # Process paragraphs
            for para in doc.paragraphs:
                text = para.text
                if not text.strip():
                    continue
                    
                conf = 1.0
                block_type = "text"
                
                # Check for MRZ in text
                mrz_result = self.mrz_detector.detect_from_text(text)
                if mrz_result:
                    block_type = "mrz"
                
                # Add block to results
                blocks.append({
                    "text": text,
                    "bbox": None,  # DOCX doesn't provide coordinates
                    "conf": conf,
                    "type": block_type,
                    "metadata": {
                        "style": para.style.name,
                        "mrz_data": mrz_result
                    } if mrz_result else {"style": para.style.name}
                })
            
            # Process tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            text = para.text
                            if not text.strip():
                                continue
                                
                            conf = 1.0
                            block_type = "text"
                            
                            # Check for MRZ in text
                            mrz_result = self.mrz_detector.detect_from_text(text)
                            if mrz_result:
                                block_type = "mrz"
                            
                            # Add block to results
                            blocks.append({
                                "text": text,
                                "bbox": None,  # DOCX doesn't provide coordinates
                                "conf": conf,
                                "type": block_type,
                                "metadata": {
                                    "content_type": "table_cell",
                                    "mrz_data": mrz_result
                                } if mrz_result else {"content_type": "table_cell"}
                            })
            
            # Add page results
            results.append({
                "page": 1,  # DOCX doesn't have page concept in API
                "text_blocks": blocks
            })
            
        except Exception as e:
            logger.error(f"DOCX processing error: {e}")
            raise
            
        return results
    
    def _extract_from_xlsx(self, file_path: str) -> List[Dict[str, Any]]:
        """
        Extract text from XLSX files
        """
        results = []
        try:
            # Open XLSX workbook
            wb = openpyxl.load_workbook(file_path, data_only=True)
            
            # Process each worksheet
            for sheet_index, sheet in enumerate(wb.worksheets):
                blocks = []
                
                # Process cells
                for row in sheet.iter_rows():
                    for cell in row:
                        # Skip empty cells
                        if cell.value is None:
                            continue
                            
                        text = str(cell.value)
                        conf = 1.0
                        block_type = "text"
                        
                        # Check for MRZ in text
                        mrz_result = self.mrz_detector.detect_from_text(text)
                        if mrz_result:
                            block_type = "mrz"
                        
                        # Add block to results
                        blocks.append({
                            "text": text,
                            "bbox": None,  # XLSX doesn't provide coordinates
                            "conf": conf,
                            "type": block_type,
                            "metadata": {
                                "sheet": sheet.title,
                                "cell": f"{cell.column_letter}{cell.row}",
                                "mrz_data": mrz_result
                            } if mrz_result else {
                                "sheet": sheet.title,
                                "cell": f"{cell.column_letter}{cell.row}"
                            }
                        })
                
                # Add sheet results
                results.append({
                    "page": sheet_index + 1,
                    "text_blocks": blocks,
                    "page_info": {
                        "sheet_name": sheet.title,
                        "max_row": sheet.max_row,
                        "max_column": sheet.max_column
                    }
                })
                
        except Exception as e:
            logger.error(f"XLSX processing error: {e}")
            raise
            
        return results
    
    def _extract_from_image(self, file_path: str) -> List[Dict[str, Any]]:
        """
        Extract text from image files
        """
        results = []
        try:
            # Open image
            img = Image.open(file_path)
            
            # Get image info
            width, height = img.size
            page_info = {"width": width, "height": height, "format": img.format}
            
            # Apply OCR
            blocks = self._apply_ocr(img)
            
            # Add page results
            results.append({
                "page": 1,
                "text_blocks": blocks,
                "page_info": page_info
            })
            
        except Exception as e:
            logger.error(f"Image processing error: {e}")
            raise
            
        return results
        
    async def extract_text(self, file_path: str, file_type: str = None) -> List[Dict[str, Any]]:
        """
        Extract text from document based on file type
        """
        # Determine file type if not provided
        if file_type is None:
            file_extension = os.path.splitext(file_path)[1].lower().lstrip('.')
            file_type = file_extension
            
        # Convert file type to lowercase for case-insensitive comparison
        file_type = file_type.lower()
        
        # Check if file type is supported
        if file_type not in SUPPORTED_TYPES:
            raise ValueError(f"Unsupported file type: {file_type}")
            
        try:
            # Process file based on type
            if file_type == "pdf":
                return self._extract_from_pdf(file_path)
                
            elif file_type == "docx":
                return self._extract_from_docx(file_path)
                
            elif file_type == "xlsx":
                return self._extract_from_xlsx(file_path)
                
            elif file_type in ["jpg", "jpeg", "png", "tiff", "tif", "bmp"]:
                return self._extract_from_image(file_path)
                
            else:
                raise ValueError(f"Unsupported file type: {file_type}")
                
        except Exception as e:
            logger.error(f"Text extraction error for {file_path}: {e}")
            raise

# Extractor main function to maintain backwards compatibility
async def extract_text(file_path, file_type=None):
    """
    Extract text from document (maintains backwards compatibility)
    """
    extractor = DocumentExtractor()
    return await extractor.extract_text(file_path, file_type)
