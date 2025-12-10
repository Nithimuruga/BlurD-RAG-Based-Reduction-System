"""
Output formatting service for redacted content.
Supports multiple output formats including PDF, DOCX, JSON, and API responses.
"""

from typing import Dict, Any, List, Optional, Union, BinaryIO, Tuple
import os
import json
import logging
import io
import base64
from pathlib import Path
from datetime import datetime
import asyncio
import uuid

# Import document manipulation libraries
import PyPDF2
from docx import Document
import pandas as pd
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter

from app.schemas.pii_schemas import PIIDetectionResult, RedactionResult, AuditLogEntry

logger = logging.getLogger(__name__)

class OutputFormatter:
    """
    Service for formatting and saving redacted content in various formats
    """
    
    def __init__(self, output_dir: str = "redacted_outputs"):
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(exist_ok=True)
    
    async def format_output(self, 
                           redaction_result: Union[RedactionResult, Dict[str, Any]],
                           format_type: str,
                           output_path: str = None,
                           metadata: Dict[str, Any] = None) -> Dict[str, Any]:
        """
        Format redacted content and save to the specified format
        
        Args:
            redaction_result: Redaction result object or dictionary
            format_type: Output format (pdf, docx, json, text, api_response)
            output_path: Optional path to save output (if None, saves to output_dir)
            metadata: Additional metadata to include in the output
            
        Returns:
            Dictionary with output information and paths
        """
        if isinstance(redaction_result, dict):
            redacted_text = redaction_result.get("redacted_text", "")
            detection_entities = redaction_result.get("detected_entities", [])
        else:
            redacted_text = redaction_result.redacted_text
            detection_entities = redaction_result.detected_entities
            
        metadata = metadata or {}
        
        # Generate a filename if not provided
        if not output_path:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"redacted_{timestamp}_{uuid.uuid4().hex[:8]}.{format_type}"
            output_path = str(self.output_dir / filename)
        
        try:
            # Process based on format type
            if format_type.lower() == "pdf":
                result = await self._create_pdf_output(redacted_text, detection_entities, output_path, metadata)
            elif format_type.lower() == "docx":
                result = await self._create_docx_output(redacted_text, detection_entities, output_path, metadata)
            elif format_type.lower() == "json":
                result = await self._create_json_output(redaction_result, output_path, metadata)
            elif format_type.lower() == "csv":
                result = await self._create_csv_output(redaction_result, output_path, metadata)
            elif format_type.lower() == "text":
                result = await self._create_text_output(redacted_text, output_path, metadata)
            elif format_type.lower() == "api_response":
                result = await self._create_api_response(redaction_result, metadata)
            else:
                return {
                    "success": False,
                    "error": f"Unsupported format type: {format_type}"
                }
                
            # Create audit log entry
            audit_entry = self._create_audit_log(redaction_result, format_type, output_path, metadata)
            result["audit"] = audit_entry
            
            return result
            
        except Exception as e:
            logger.error(f"Error formatting output: {e}")
            return {
                "success": False,
                "error": str(e)
            }
    
    async def _create_pdf_output(self, 
                                redacted_text: str, 
                                detection_entities: List[Any],
                                output_path: str, 
                                metadata: Dict[str, Any]) -> Dict[str, Any]:
        """Create a PDF file with redacted content"""
        try:
            # Check if we're redacting an existing PDF
            original_pdf_path = metadata.get("original_file_path")
            if original_pdf_path and original_pdf_path.lower().endswith(".pdf") and os.path.exists(original_pdf_path):
                # Modify existing PDF
                return await self._redact_existing_pdf(original_pdf_path, redacted_text, detection_entities, output_path)
            else:
                # Create new PDF from scratch
                buffer = io.BytesIO()
                c = canvas.Canvas(buffer, pagesize=letter)
                width, height = letter
                
                # Add header with metadata
                c.setFont("Helvetica-Bold", 14)
                c.drawString(72, height - 72, "Redacted Document")
                
                c.setFont("Helvetica", 10)
                c.drawString(72, height - 90, f"Processed: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
                
                if metadata.get("document_id"):
                    c.drawString(72, height - 108, f"Document ID: {metadata['document_id']}")
                
                # Add the redacted text
                c.setFont("Helvetica", 10)
                text_obj = c.beginText(72, height - 144)
                
                # Simple word wrapping
                words = redacted_text.split()
                line = ""
                for word in words:
                    if len(line + " " + word) * 6 < width - 144:  # Approximate width
                        line = line + " " + word if line else word
                    else:
                        text_obj.textLine(line)
                        line = word
                
                if line:
                    text_obj.textLine(line)
                
                c.drawText(text_obj)
                
                # Add footer with redaction statistics
                c.setFont("Helvetica-Oblique", 9)
                
                # Count entities by type
                type_counts = {}
                for entity in detection_entities:
                    entity_type = entity.pii_type if hasattr(entity, "pii_type") else entity.get("pii_type", "unknown")
                    type_counts[entity_type] = type_counts.get(entity_type, 0) + 1
                
                footer_text = "Redaction summary: "
                for entity_type, count in type_counts.items():
                    footer_text += f"{entity_type}={count}, "
                
                footer_text = footer_text.rstrip(", ")
                c.drawString(72, 36, footer_text)
                
                # Finalize PDF
                c.save()
                
                # Save the PDF to file
                with open(output_path, "wb") as f:
                    f.write(buffer.getvalue())
                
                return {
                    "success": True,
                    "format": "pdf",
                    "output_path": output_path,
                    "size_bytes": os.path.getsize(output_path)
                }
        
        except Exception as e:
            logger.error(f"Error creating PDF output: {e}")
            return {
                "success": False,
                "format": "pdf",
                "error": str(e)
            }
    
    async def _redact_existing_pdf(self, 
                                  original_pdf: str,
                                  redacted_text: str, 
                                  detection_entities: List[Any],
                                  output_path: str) -> Dict[str, Any]:
        """Redact content in an existing PDF file"""
        # Note: Full PDF redaction with proper visual overlay requires more sophisticated libraries
        # This is a simplified version that creates a new PDF with the redacted text
        try:
            buffer = io.BytesIO()
            c = canvas.Canvas(buffer, pagesize=letter)
            width, height = letter
            
            # Add header
            c.setFont("Helvetica-Bold", 14)
            c.drawString(72, height - 72, "Redacted Document (Original PDF)")
            
            c.setFont("Helvetica", 10)
            c.drawString(72, height - 90, f"Processed: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            c.drawString(72, height - 108, f"Original PDF: {os.path.basename(original_pdf)}")
            
            # Add the redacted text
            c.setFont("Helvetica", 10)
            text_obj = c.beginText(72, height - 144)
            
            # Simple word wrapping
            words = redacted_text.split()
            line = ""
            for word in words:
                if len(line + " " + word) * 6 < width - 144:  # Approximate width
                    line = line + " " + word if line else word
                else:
                    text_obj.textLine(line)
                    line = word
            
            if line:
                text_obj.textLine(line)
            
            c.drawText(text_obj)
            
            # Finalize PDF
            c.save()
            
            # Save the PDF to file
            with open(output_path, "wb") as f:
                f.write(buffer.getvalue())
            
            return {
                "success": True,
                "format": "pdf",
                "output_path": output_path,
                "size_bytes": os.path.getsize(output_path),
                "note": "Simple text replacement. For full PDF redaction with visual overlay, advanced PDF libraries are required."
            }
            
        except Exception as e:
            logger.error(f"Error redacting existing PDF: {e}")
            return {
                "success": False,
                "format": "pdf",
                "error": str(e)
            }
    
    async def _create_docx_output(self, 
                                 redacted_text: str, 
                                 detection_entities: List[Any],
                                 output_path: str, 
                                 metadata: Dict[str, Any]) -> Dict[str, Any]:
        """Create a DOCX file with redacted content"""
        try:
            doc = Document()
            
            # Add title
            doc.add_heading("Redacted Document", level=1)
            
            # Add metadata
            doc.add_paragraph(f"Processed: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            
            if metadata.get("document_id"):
                doc.add_paragraph(f"Document ID: {metadata['document_id']}")
            
            # Add horizontal rule
            doc.add_paragraph("_" * 50)
            
            # Add redacted content
            doc.add_paragraph(redacted_text)
            
            # Add redaction statistics
            doc.add_heading("Redaction Statistics", level=2)
            
            # Count entities by type
            type_counts = {}
            for entity in detection_entities:
                entity_type = entity.pii_type if hasattr(entity, "pii_type") else entity.get("pii_type", "unknown")
                type_counts[entity_type] = type_counts.get(entity_type, 0) + 1
            
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            
            # Add header row
            header_cells = table.rows[0].cells
            header_cells[0].text = "PII Type"
            header_cells[1].text = "Count"
            
            # Add data rows
            for entity_type, count in type_counts.items():
                row_cells = table.add_row().cells
                row_cells[0].text = str(entity_type)
                row_cells[1].text = str(count)
            
            # Save the document
            doc.save(output_path)
            
            return {
                "success": True,
                "format": "docx",
                "output_path": output_path,
                "size_bytes": os.path.getsize(output_path)
            }
            
        except Exception as e:
            logger.error(f"Error creating DOCX output: {e}")
            return {
                "success": False,
                "format": "docx",
                "error": str(e)
            }
    
    async def _create_json_output(self, 
                                 redaction_result: Union[RedactionResult, Dict[str, Any]],
                                 output_path: str, 
                                 metadata: Dict[str, Any]) -> Dict[str, Any]:
        """Create a JSON file with redacted content and metadata"""
        try:
            # Convert to dictionary if it's an object
            if not isinstance(redaction_result, dict):
                result_dict = {
                    "redacted_text": redaction_result.redacted_text,
                    "redaction_count": redaction_result.redaction_count,
                    "processed_at": datetime.now().isoformat(),
                    "metadata": metadata
                }
                
                # Convert entities to dictionaries
                entities = []
                for entity in redaction_result.detected_entities:
                    if hasattr(entity, "dict"):
                        entities.append(entity.dict())
                    else:
                        entities.append(dict(entity))
                        
                result_dict["detected_entities"] = entities
            else:
                result_dict = {
                    **redaction_result,
                    "processed_at": datetime.now().isoformat(),
                    "metadata": metadata
                }
            
            # Write to JSON file
            with open(output_path, "w", encoding="utf-8") as f:
                json.dump(result_dict, f, indent=2, ensure_ascii=False)
            
            return {
                "success": True,
                "format": "json",
                "output_path": output_path,
                "size_bytes": os.path.getsize(output_path)
            }
            
        except Exception as e:
            logger.error(f"Error creating JSON output: {e}")
            return {
                "success": False,
                "format": "json",
                "error": str(e)
            }
    
    async def _create_csv_output(self, 
                               redaction_result: Union[RedactionResult, Dict[str, Any]],
                               output_path: str, 
                               metadata: Dict[str, Any]) -> Dict[str, Any]:
        """Create a CSV file with detected entities"""
        try:
            # Extract entities
            if isinstance(redaction_result, dict):
                entities = redaction_result.get("detected_entities", [])
            else:
                entities = redaction_result.detected_entities
            
            # Prepare data for CSV
            rows = []
            for entity in entities:
                if hasattr(entity, "pii_type"):
                    # It's an object
                    row = {
                        "pii_type": str(entity.pii_type),
                        "original_text": entity.text,
                        "redacted_text": entity.redacted_text or "",
                        "confidence": entity.confidence,
                        "risk_level": str(entity.risk_level),
                        "start_position": entity.start_position,
                        "end_position": entity.end_position,
                    }
                else:
                    # It's a dictionary
                    row = {
                        "pii_type": entity.get("pii_type", "unknown"),
                        "original_text": entity.get("text", ""),
                        "redacted_text": entity.get("redacted_text", ""),
                        "confidence": entity.get("confidence", 0),
                        "risk_level": entity.get("risk_level", "unknown"),
                        "start_position": entity.get("start_position", 0),
                        "end_position": entity.get("end_position", 0),
                    }
                
                rows.append(row)
            
            # Create DataFrame and save to CSV
            df = pd.DataFrame(rows)
            df.to_csv(output_path, index=False)
            
            return {
                "success": True,
                "format": "csv",
                "output_path": output_path,
                "size_bytes": os.path.getsize(output_path),
                "row_count": len(rows)
            }
            
        except Exception as e:
            logger.error(f"Error creating CSV output: {e}")
            return {
                "success": False,
                "format": "csv",
                "error": str(e)
            }
    
    async def _create_text_output(self, 
                                redacted_text: str,
                                output_path: str, 
                                metadata: Dict[str, Any]) -> Dict[str, Any]:
        """Create a plain text file with redacted content"""
        try:
            # Add a simple header
            header = "===== REDACTED DOCUMENT =====\n"
            header += f"Processed: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n"
            
            if metadata.get("document_id"):
                header += f"Document ID: {metadata['document_id']}\n"
                
            header += "=" * 30 + "\n\n"
            
            # Combine header and content
            content = header + redacted_text
            
            # Write to file
            with open(output_path, "w", encoding="utf-8") as f:
                f.write(content)
            
            return {
                "success": True,
                "format": "text",
                "output_path": output_path,
                "size_bytes": os.path.getsize(output_path)
            }
            
        except Exception as e:
            logger.error(f"Error creating text output: {e}")
            return {
                "success": False,
                "format": "text",
                "error": str(e)
            }
    
    async def _create_api_response(self, 
                                 redaction_result: Union[RedactionResult, Dict[str, Any]],
                                 metadata: Dict[str, Any]) -> Dict[str, Any]:
        """Format redaction result for API response"""
        try:
            # Extract necessary data
            if isinstance(redaction_result, dict):
                redacted_text = redaction_result.get("redacted_text", "")
                
                # Make a copy to avoid modifying the original
                response = {
                    "success": True,
                    "redacted_text": redacted_text,
                    "metadata": metadata,
                    "timestamp": datetime.now().isoformat()
                }
                
                # Include redaction counts if available
                if "redaction_count" in redaction_result:
                    response["redaction_statistics"] = {
                        "total_redactions": sum(redaction_result["redaction_count"].values()),
                        "by_type": redaction_result["redaction_count"]
                    }
                    
                # Include processing time if available
                if "processing_time" in redaction_result:
                    response["processing_time"] = redaction_result["processing_time"]
            else:
                # It's a RedactionResult object
                response = {
                    "success": True,
                    "redacted_text": redaction_result.redacted_text,
                    "metadata": metadata,
                    "timestamp": datetime.now().isoformat(),
                    "redaction_statistics": {
                        "total_redactions": sum(redaction_result.redaction_count.values()),
                        "by_type": redaction_result.redaction_count
                    },
                    "processing_time": redaction_result.processing_time
                }
            
            return response
            
        except Exception as e:
            logger.error(f"Error creating API response: {e}")
            return {
                "success": False,
                "error": str(e)
            }
    
    def _create_audit_log(self,
                         redaction_result: Union[RedactionResult, Dict[str, Any]],
                         format_type: str,
                         output_path: str,
                         metadata: Dict[str, Any]) -> Dict[str, Any]:
        """Create an audit log entry for the redaction operation"""
        try:
            # Extract user ID if available
            user_id = metadata.get("user_id")
            document_id = metadata.get("document_id") or metadata.get("file_id")
            source_ip = metadata.get("source_ip")
            
            # Get PII types processed
            if isinstance(redaction_result, dict):
                pii_types_dict = redaction_result.get("redaction_count", {})
                entity_count = sum(pii_types_dict.values())
            else:
                pii_types_dict = redaction_result.redaction_count
                entity_count = sum(pii_types_dict.values())
                
            pii_types = list(pii_types_dict.keys())
            
            # Create audit log entry
            audit_entry = AuditLogEntry(
                timestamp=datetime.utcnow(),
                operation=f"redact_{format_type}",
                user_id=user_id,
                document_id=document_id,
                pii_types_processed=pii_types,
                entity_count=entity_count,
                success=True,
                compliance_frameworks=metadata.get("compliance_frameworks", []),
                source_ip=source_ip,
                metadata={
                    "output_path": output_path,
                    "format": format_type
                }
            )
            
            # Convert to dictionary for return
            return audit_entry.dict()
            
        except Exception as e:
            logger.error(f"Error creating audit log: {e}")
            return {
                "timestamp": datetime.utcnow().isoformat(),
                "operation": f"redact_{format_type}",
                "success": False,
                "error": str(e)
            }